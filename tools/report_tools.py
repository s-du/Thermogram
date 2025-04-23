"""
Report generation tools for DroneIR Toolkit.
- Automated Word report creation with sections:
  Contents / Objectives / Site and Conditions / Flight details / Images
- Images section is filled with selected images (ProcessedIm)
- Embedded annotations summarized as tables
"""
import os
import io
import tempfile
import numpy as np
from PIL import Image
import cv2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import _Cell
from .core import ProcessedIm, cv_read_all_path

STYLE_TEMPLATES = {
    "modern_blue": {
        # Text styles
        "Title": {"font": "Arial", "size": 28, "bold": True, "color": (0x2E, 0x74, 0xB5), "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER},
        "Subtitle": {"font": "Arial", "size": 16, "italic": True, "color": (0x4F, 0x81, 0xBD), "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER},
        "Heading 1": {"font": "Arial", "size": 18, "bold": True, "color": (0x2E, 0x74, 0xB5), "spacing_before": 12, "spacing_after": 6},
        "Heading 2": {"font": "Arial", "size": 15, "bold": True, "color": (0x4F, 0x81, 0xBD), "spacing_before": 10, "spacing_after": 4},
        "Heading 3": {"font": "Arial", "size": 13, "bold": True, "color": (0x1F, 0x49, 0x7D), "spacing_before": 8, "spacing_after": 2},
        "Normal": {"font": "Calibri", "size": 11, "color": (0, 0, 0), "line_spacing": 1.15},
        "Caption": {"font": "Calibri", "size": 10, "italic": True, "color": (0x4F, 0x81, 0xBD)},
        "Table Grid": {"font": "Calibri", "size": 10, "color": (0, 0, 0), "shading": (0xF2, 0xF9, 0xFF)},
        
        # Cover page
        "cover_bg_color": None,  # No background color
        "cover_accent_color": (0x2E, 0x74, 0xB5),  # Blue accent
        "cover_layout": "centered",  # Title centered on page
        "cover_image": None,  # No default image
        "cover_footer": True,  # Include footer with date
        
        # Table styles
        "table_header_bg": (0x4F, 0x81, 0xBD),
        "table_header_text": (0xFF, 0xFF, 0xFF),
        "table_row_alt": (0xE6, 0xEF, 0xF7),
        "table_border": (0x4F, 0x81, 0xBD),
        
        # Other elements
        "page_margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},  # in inches
        "header_footer": True,  # Include header/footer
        "watermark": None,  # No watermark
    },
    
    "classic_gray": {
        # Text styles
        "Title": {"font": "Times New Roman", "size": 30, "bold": True, "color": (0x33, 0x33, 0x33), "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER},
        "Subtitle": {"font": "Times New Roman", "size": 18, "italic": True, "color": (0x66, 0x66, 0x66), "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER},
        "Heading 1": {"font": "Times New Roman", "size": 20, "bold": True, "color": (0x44, 0x44, 0x44), "spacing_before": 14, "spacing_after": 7},
        "Heading 2": {"font": "Times New Roman", "size": 16, "bold": True, "color": (0x66, 0x66, 0x66), "spacing_before": 12, "spacing_after": 6},
        "Heading 3": {"font": "Times New Roman", "size": 13, "bold": True, "color": (0x88, 0x88, 0x88), "spacing_before": 10, "spacing_after": 5},
        "Normal": {"font": "Times New Roman", "size": 12, "color": (0x33, 0x33, 0x33), "line_spacing": 1.08},
        "Caption": {"font": "Times New Roman", "size": 11, "italic": True, "color": (0x66, 0x66, 0x66)},
        "Table Grid": {"font": "Times New Roman", "size": 10, "color": (0x44, 0x44, 0x44), "shading": (0xF6, 0xF6, 0xF6)},
        
        # Cover page
        "cover_bg_color": (0xF8, 0xF8, 0xF8),  # Light gray background
        "cover_accent_color": (0x88, 0x88, 0x88),  # Gray accent
        "cover_layout": "traditional",  # Title at top with spacing
        "cover_image": None,  # No default image
        "cover_footer": True,  # Include footer with date
        
        # Table styles
        "table_header_bg": (0x88, 0x88, 0x88),
        "table_header_text": (0xFF, 0xFF, 0xFF),
        "table_row_alt": (0xF0, 0xF0, 0xF0),
        "table_border": (0xAA, 0xAA, 0xAA),
        
        # Other elements
        "page_margins": {"top": 1.25, "bottom": 1.25, "left": 1.25, "right": 1.25},  # in inches
        "header_footer": True,  # Include header/footer
        "watermark": None,  # No watermark
    },
    
    "dark_elegant": {
        # Text styles
        "Title": {"font": "Segoe UI", "size": 32, "bold": True, "color": (0xCC, 0xCC, 0xCC), "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER},
        "Subtitle": {"font": "Segoe UI Light", "size": 18, "italic": False, "color": (0xAA, 0xAA, 0xAA), "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER},
        "Heading 1": {"font": "Segoe UI", "size": 18, "bold": True, "color": (0xAA, 0xAA, 0xAA), "spacing_before": 14, "spacing_after": 7},
        "Heading 2": {"font": "Segoe UI", "size": 15, "bold": True, "color": (0xBB, 0xBB, 0xBB), "spacing_before": 12, "spacing_after": 6},
        "Heading 3": {"font": "Segoe UI", "size": 13, "bold": True, "color": (0xCC, 0xCC, 0xCC), "spacing_before": 10, "spacing_after": 5},
        "Normal": {"font": "Segoe UI", "size": 11, "color": (0x22, 0x22, 0x22), "line_spacing": 1.2},
        "Caption": {"font": "Segoe UI Light", "size": 10, "italic": True, "color": (0x88, 0x88, 0x88)},
        "Table Grid": {"font": "Segoe UI", "size": 10, "color": (0x22, 0x22, 0x22), "shading": (0xE5, 0xE5, 0xE5)},
        
        # Cover page
        "cover_bg_color": (0x22, 0x22, 0x22),  # Dark background
        "cover_accent_color": (0xE5, 0xE5, 0xE5),  # Light accent
        "cover_layout": "modern",  # Title offset to left with accent line
        "cover_image": None,  # No default image
        "cover_footer": True,  # Include footer with date
        
        # Table styles
        "table_header_bg": (0x44, 0x44, 0x44),
        "table_header_text": (0xE5, 0xE5, 0xE5),
        "table_row_alt": (0xF5, 0xF5, 0xF5),
        "table_border": (0x66, 0x66, 0x66),
        
        # Other elements
        "page_margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},  # in inches
        "header_footer": True,  # Include header/footer
        "watermark": None,  # No watermark
    },
    
    "vibrant_thermal": {
        # Text styles
        "Title": {"font": "Segoe UI", "size": 32, "bold": True, "color": (0xFF, 0x45, 0x00), "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER},  # Thermal orange
        "Subtitle": {"font": "Segoe UI Light", "size": 18, "italic": False, "color": (0xFF, 0x8C, 0x00), "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER},  # Orange
        "Heading 1": {"font": "Segoe UI", "size": 18, "bold": True, "color": (0xFF, 0x45, 0x00), "spacing_before": 14, "spacing_after": 7},  # Thermal orange
        "Heading 2": {"font": "Segoe UI", "size": 15, "bold": True, "color": (0xFF, 0x8C, 0x00), "spacing_before": 12, "spacing_after": 6},  # Orange
        "Heading 3": {"font": "Segoe UI", "size": 13, "bold": True, "color": (0xCC, 0x00, 0x00), "spacing_before": 10, "spacing_after": 5},  # Dark red
        "Normal": {"font": "Calibri", "size": 11, "color": (0x33, 0x33, 0x33), "line_spacing": 1.15},
        "Caption": {"font": "Calibri", "size": 10, "italic": True, "color": (0xFF, 0x8C, 0x00)},  # Orange
        "Table Grid": {"font": "Calibri", "size": 10, "color": (0x33, 0x33, 0x33), "shading": (0xFF, 0xF5, 0xE6)},  # Light orange
        
        # Cover page
        "cover_bg_color": (0x33, 0x33, 0x33),  # Dark gray background
        "cover_accent_color": (0xFF, 0x45, 0x00),  # Thermal orange accent
        "cover_layout": "gradient",  # Gradient background
        "cover_image": None,  # No default image
        "cover_footer": True,  # Include footer with date
        
        # Table styles
        "table_header_bg": (0xFF, 0x45, 0x00),  # Thermal orange
        "table_header_text": (0xFF, 0xFF, 0xFF),  # White
        "table_row_alt": (0xFF, 0xF5, 0xE6),  # Light orange
        "table_border": (0xFF, 0x8C, 0x00),  # Orange
        
        # Other elements
        "page_margins": {"top": 0.75, "bottom": 0.75, "left": 0.75, "right": 0.75},  # in inches
        "header_footer": True,  # Include header/footer
        "watermark": None,  # No watermark
    }
}

def set_custom_styles(doc, style_template="modern_blue"):
    """
    Apply the chosen style template to the Word document.
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.enum.section import WD_SECTION, WD_ORIENT
    
    # Get the style template dictionary
    if isinstance(style_template, str):
        tpl = STYLE_TEMPLATES.get(style_template, STYLE_TEMPLATES["modern_blue"])
    else:
        tpl = style_template
    
    # Set page margins if specified
    if "page_margins" in tpl:
        margins = tpl["page_margins"]
        section = doc.sections[0]
        section.top_margin = Inches(margins.get("top", 1.0))
        section.bottom_margin = Inches(margins.get("bottom", 1.0))
        section.left_margin = Inches(margins.get("left", 1.0))
        section.right_margin = Inches(margins.get("right", 1.0))
    
    # Apply styles to document
    for style_name, style_props in tpl.items():
        if style_name in doc.styles:
            style = doc.styles[style_name]
            
            # Set font properties
            if "font" in style_props:
                style.font.name = style_props["font"]
            if "size" in style_props:
                style.font.size = Pt(style_props["size"])
            if "bold" in style_props:
                style.font.bold = style_props["bold"]
            if "italic" in style_props:
                style.font.italic = style_props["italic"]
            if "color" in style_props:
                style.font.color.rgb = RGBColor(*style_props["color"])
            
            # Set paragraph properties
            if "alignment" in style_props:
                style.paragraph_format.alignment = style_props["alignment"]
            if "spacing_before" in style_props:
                style.paragraph_format.space_before = Pt(style_props["spacing_before"])
            if "spacing_after" in style_props:
                style.paragraph_format.space_after = Pt(style_props["spacing_after"])
            if "line_spacing" in style_props:
                style.paragraph_format.line_spacing = style_props["line_spacing"]
    
    return tpl


def add_table_from_annotations(doc, annotations, title="Annotations Summary", style_template=None, annotation_type=None):
    """
    Add a table summarizing annotations to the doc.
    annotations: list of dicts or objects with annotation info
    style_template: dict with table style settings (optional)
    annotation_type: type of annotation ('rect', 'point', 'line') to customize display
    """
    if not annotations:
        return
    
    doc.add_heading(title, level=3)
    
    # Convert objects to dicts if needed
    ann_dicts = []
    for ann in annotations:
        if not isinstance(ann, dict):
            ann_dict = {k: v for k, v in ann.__dict__.items() 
                      if not k.startswith('_') and not callable(v) and not isinstance(v, (list, tuple, dict)) 
                      and not any(excluded in k for excluded in ['item', 'ellipse', 'text', 'spot', 'main'])}
            ann_dicts.append(ann_dict)
        else:
            ann_dicts.append(ann)
    
    # Select relevant keys based on annotation type
    if annotation_type == 'rect':
        # For rectangular measurements, focus on temperature data and area
        keys = ['name', 'tmin', 'tmax', 'tmean', 'tstd', 'area']
        headers = ['Name', 'Min Temp (°C)', 'Max Temp (°C)', 'Mean Temp (°C)', 'Std Dev (°C)', 'Area (px²)']
    elif annotation_type == 'point':
        # For point measurements, show position and temperature
        keys = ['name', 'temp']
        headers = ['Name', 'Temperature (°C)']
    elif annotation_type == 'line':
        # For line measurements, show min/max/mean temperatures
        keys = ['name', 'tmin', 'tmax', 'tmean']
        headers = ['Name', 'Min Temp (°C)', 'Max Temp (°C)', 'Mean Temp (°C)']
    else:
        # Default: use all available keys (filtered)
        keys = []
        for ann in ann_dicts:
            for k in ann.keys():
                if k not in keys and not k.startswith('_') and not any(excluded in k for excluded in ['item', 'ellipse', 'text', 'spot', 'main', 'coords', 'data_roi']):
                    keys.append(k)
        headers = [k.replace('_', ' ').title() for k in keys]
    
    # Create table with selected keys
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Apply table shading if specified
    if style_template and "Table Grid" in style_template:
        shading = style_template["Table Grid"].get("shading")
        if shading:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            for row in table.rows:
                for cell in row.cells:
                    cell._tc.get_or_add_tcPr().append(parse_xml(
                        r'<w:shd {} w:fill="{:02X}{:02X}{:02X}"/>'.format(
                            nsdecls('w'), *shading)))
    
    # Set headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Make headers bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Fill table with data
    for ann in ann_dicts:
        row_cells = table.add_row().cells
        for i, key in enumerate(keys):
            value = ann.get(key, 'N/A')
            # Format numeric values
            if isinstance(value, float):
                row_cells[i].text = f"{value:.2f}"
            else:
                row_cells[i].text = str(value)
    
    # Apply enhanced table styling based on template
    if isinstance(style_template, str):
        tpl = STYLE_TEMPLATES.get(style_template, STYLE_TEMPLATES["modern_blue"])
        
        # Apply header styling
        if "table_header_bg" in tpl and "table_header_text" in tpl:
            header_bg = tpl["table_header_bg"]
            header_text = tpl["table_header_text"]
            
            # Apply to header row
            for cell in table.rows[0].cells:
                # Set background color
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                cell_xml = cell._tc.get_or_add_tcPr()
                cell_xml.append(parse_xml(
                    r'<w:shd {} w:fill="{:02X}{:02X}{:02X}"/>'.format(
                        nsdecls('w'), *header_bg)))
                
                # Set text color
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(*header_text)
        
        # Apply alternating row colors if specified
        if "table_row_alt" in tpl:
            alt_color = tpl["table_row_alt"]
            for i, row in enumerate(table.rows):
                # Skip header row
                if i == 0:
                    continue
                    
                # Apply to alternating rows
                if i % 2 == 1:  # Apply to odd rows (1-indexed, so 2nd, 4th, etc.)
                    for cell in row.cells:
                        from docx.oxml import parse_xml
                        from docx.oxml.ns import nsdecls
                        cell_xml = cell._tc.get_or_add_tcPr()
                        cell_xml.append(parse_xml(
                            r'<w:shd {} w:fill="{:02X}{:02X}{:02X}"/>'.format(
                                nsdecls('w'), *alt_color)))
    
    doc.add_paragraph()
    return table


def safe_add_image_to_doc(doc, image_path, width=None, height=None):
    """
    Safely add an image to a Word document, handling various image formats
    and potential issues with thermal images.
    
    Args:
        doc: The Word document object
        image_path: Path to the image file
        width: Optional width for the image
        height: Optional height for the image
    """
    try:
        # First try direct addition (fastest)
        doc.add_picture(image_path, width=width, height=height)
    except Exception:
        # If direct addition fails, try conversion with PIL
        try:
            # Open with PIL which handles more formats
            with Image.open(image_path) as img:
                # Convert to RGB if needed (some thermal images might be in unusual formats)
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                    
                # Save to a temporary file
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                    temp_path = temp_file.name
                    img.save(temp_path, format='PNG')
                
                # Add the converted image
                doc.add_picture(temp_path, width=width, height=height)
                
                # Clean up
                os.unlink(temp_path)
        except Exception:
            # If PIL fails, try with OpenCV
            try:
                # Read with OpenCV
                img = cv_read_all_path(image_path)
                if img is None:
                    raise ValueError("Could not read image with OpenCV")
                    
                # Convert BGR to RGB if needed
                if len(img.shape) == 3 and img.shape[2] == 3:
                    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
                    
                # Save to a temporary file
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                    temp_path = temp_file.name
                    cv2.imwrite(temp_path, img)
                
                # Add the converted image
                doc.add_picture(temp_path, width=width, height=height)
                
                # Clean up
                os.unlink(temp_path)
            except Exception as e:
                # If all methods fail, re-raise the exception
                raise ValueError(f"Could not process image: {str(e)}")


def create_word_report(
    output_path,
    objectives_text,
    site_conditions_text,
    flight_details_text,
    processed_images,
    images_to_include=None,
    style_template="modern_blue",
    include_summary=True,
    annotated_image_paths=None,
    report_title="Infrared Survey Report",
    report_subtitle=None,
    cover_image_path=None
):
    """
    Create a Word report with specified sections and images.
    processed_images: list of ProcessedIm objects
    images_to_include: list of indices or ProcessedIm objects to include (default: all)
    style_template: string key for the style template to use
    report_title: Title for the report cover page
    report_subtitle: Optional subtitle for the report cover page
    cover_image_path: Optional path to an image for the cover page
    """
    from datetime import datetime
    
    # Create document and apply styles
    doc = Document()
    tpl = set_custom_styles(doc, style_template)
    
    # Create cover page based on template
    cover_layout = tpl.get("cover_layout", "centered")
    cover_bg_color = tpl.get("cover_bg_color")
    cover_accent_color = tpl.get("cover_accent_color")
    
    # Add a section break for the cover page
    doc.add_section()
    
    # Apply background color to cover page if specified
    if cover_bg_color:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        section_props = doc.sections[0]._sectPr
        bg_color_xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{cover_bg_color[0]:02X}{cover_bg_color[1]:02X}{cover_bg_color[2]:02X}"/>')
        section_props.append(bg_color_xml)
    
    # Add cover image if provided
    if cover_image_path and os.path.exists(cover_image_path):
        doc.add_picture(cover_image_path, width=Inches(6.0))
    elif tpl.get("cover_image"):
        # Use template default image if available
        doc.add_picture(tpl["cover_image"], width=Inches(6.0))
    
    # Add spacing before title
    if cover_layout == "traditional":
        for _ in range(4):  # Add more space at top for traditional layout
            doc.add_paragraph()
    elif cover_layout == "modern":
        doc.add_paragraph()  # Less space for modern layout
        # Add accent line for modern layout
        if cover_accent_color:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(1.0)
            accent_run = p.add_run("■■■■■■■■■")
            accent_run.font.color.rgb = RGBColor(*cover_accent_color)
            accent_run.font.size = Pt(14)
            accent_run.bold = True
    
    # Add title with style from template
    title_style = tpl.get("Title", {})
    title_p = doc.add_paragraph()
    if "alignment" in title_style:
        title_p.paragraph_format.alignment = title_style["alignment"]
    else:
        title_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    title_run = title_p.add_run(report_title)
    title_run.bold = title_style.get("bold", True)
    title_run.font.size = Pt(title_style.get("size", 28))
    title_run.font.name = title_style.get("font", "Arial")
    if "color" in title_style:
        title_run.font.color.rgb = RGBColor(*title_style["color"])
    
    # Add subtitle if provided
    if report_subtitle:
        subtitle_style = tpl.get("Subtitle", {})
        subtitle_p = doc.add_paragraph()
        if "alignment" in subtitle_style:
            subtitle_p.paragraph_format.alignment = subtitle_style["alignment"]
        else:
            subtitle_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        subtitle_run = subtitle_p.add_run(report_subtitle)
        subtitle_run.italic = subtitle_style.get("italic", True)
        subtitle_run.bold = subtitle_style.get("bold", False)
        subtitle_run.font.size = Pt(subtitle_style.get("size", 16))
        subtitle_run.font.name = subtitle_style.get("font", "Arial")
        if "color" in subtitle_style:
            subtitle_run.font.color.rgb = RGBColor(*subtitle_style["color"])
    
    # Add date to cover page footer if specified
    if tpl.get("cover_footer", True):
        # Add some space
        for _ in range(6 if cover_layout == "traditional" else 4):
            doc.add_paragraph()
        
        date_p = doc.add_paragraph()
        date_p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        date_run = date_p.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(10)
        date_run.font.italic = True
        if cover_accent_color:
            date_run.font.color.rgb = RGBColor(*cover_accent_color)
    
    # Add page break after cover
    doc.add_page_break()
    
    # Contents
    doc.add_heading('Contents', level=1)
    toc = doc.add_paragraph()
    toc.add_run('1. Objectives\n').bold = True
    toc.add_run('2. Site and Conditions\n').bold = True
    toc.add_run('3. Flight details\n').bold = True
    toc.add_run('4. Images\n').bold = True
    doc.add_page_break()

    # Objectives
    doc.add_heading('1. Objectives', level=1)
    doc.add_paragraph(objectives_text)
    doc.add_page_break()

    # Site and Conditions
    doc.add_heading('2. Site and Conditions', level=1)
    doc.add_paragraph(site_conditions_text)
    doc.add_page_break()

    # Flight details
    doc.add_heading('3. Flight details', level=1)
    doc.add_paragraph(flight_details_text)
    doc.add_page_break()

    # Images
    doc.add_heading('4. Images', level=1)
    
    # Add summary section if requested
    if include_summary and processed_images:
        doc.add_heading('Summary of Findings', level=2)
        summary_para = doc.add_paragraph()
        summary_para.add_run('This section provides a summary of key thermal findings across all analyzed images.')
        
        # Create summary table of key measurements
        all_rect_measurements = []
        for img_idx, img in enumerate(processed_images):
            pim = img if isinstance(img, ProcessedIm) else img
            if hasattr(pim, 'meas_rect_list') and pim.meas_rect_list:
                for rect in pim.meas_rect_list:
                    if hasattr(rect, 'tmax') and hasattr(rect, 'tmin'):
                        name = getattr(rect, 'name', f'Area in Image {img_idx}')
                        all_rect_measurements.append({
                            'Image': img_idx,
                            'Name': name,
                            'Min Temp': getattr(rect, 'tmin', 'N/A'),
                            'Max Temp': getattr(rect, 'tmax', 'N/A'),
                            'Delta T': getattr(rect, 'tmax', 0) - getattr(rect, 'tmin', 0) if hasattr(rect, 'tmax') and hasattr(rect, 'tmin') else 'N/A'
                        })
        
        if all_rect_measurements:
            # Sort by max temperature to highlight potential issues
            all_rect_measurements.sort(key=lambda x: x['Max Temp'] if isinstance(x['Max Temp'], (int, float)) else -999, reverse=True)
            
            # Add summary table
            doc.add_heading('Temperature Hotspots', level=3)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Headers
            headers = ['Image', 'Name', 'Min Temp (°C)', 'Max Temp (°C)', 'Delta T (°C)'] 
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                for paragraph in hdr_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Data rows (limit to top 5 for clarity)
            for measurement in all_rect_measurements[:5]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(measurement['Image'])
                row_cells[1].text = str(measurement['Name'])
                row_cells[2].text = f"{measurement['Min Temp']:.2f}" if isinstance(measurement['Min Temp'], (int, float)) else str(measurement['Min Temp'])
                row_cells[3].text = f"{measurement['Max Temp']:.2f}" if isinstance(measurement['Max Temp'], (int, float)) else str(measurement['Max Temp'])
                row_cells[4].text = f"{measurement['Delta T']:.2f}" if isinstance(measurement['Delta T'], (int, float)) else str(measurement['Delta T'])
            
            doc.add_paragraph()
        
        doc.add_page_break()
    
    # Individual images
    if images_to_include is None:
        images_to_include = list(range(len(processed_images)))
    for idx in images_to_include:
        pim = processed_images[idx] if isinstance(processed_images[idx], ProcessedIm) else processed_images[idx]
        doc.add_heading(f'Image {idx}', level=2)
        # Add annotated thermal image with measurements if available
        if annotated_image_paths and idx < len(annotated_image_paths) and os.path.exists(annotated_image_paths[idx]):
            doc.add_heading('Thermal Image with Annotations', level=3)
            try:
                # Try to safely add the annotated thermal image
                safe_add_image_to_doc(doc, annotated_image_paths[idx], width=Inches(6))
            except Exception as e:
                # If annotated image fails, add a note
                doc.add_paragraph(f"[Annotated image could not be displayed: {str(e)}]", style='Normal')
        
        # Add thermal image if available (raw thermal image without annotations)
        if pim.path and os.path.exists(pim.path):
            doc.add_heading('Raw Thermal Image', level=3)
            try:
                # Try to safely add the thermal image using PIL to convert if needed
                safe_add_image_to_doc(doc, pim.path, width=Inches(4))
            except Exception as e:
                # If thermal image fails, add a note
                doc.add_paragraph(f"[Thermal image could not be displayed: {str(e)}]", style='Normal')
            
        # Add RGB image if available
        if pim.rgb_path and os.path.exists(pim.rgb_path):
            doc.add_heading('RGB Image', level=3)
            try:
                # Try to safely add the RGB image
                safe_add_image_to_doc(doc, pim.rgb_path, width=Inches(4))
            except Exception as e:
                # If RGB image fails, add a note
                doc.add_paragraph(f"[RGB image could not be displayed: {str(e)}]", style='Normal')
        # Add image info section
        doc.add_heading('Image Information', level=3)
        
        # Add basic image info
        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = 'Table Grid'
        info_table.autofit = True
        
        # Set column widths
        info_table.columns[0].width = Inches(1.5)
        info_table.columns[1].width = Inches(4.5)
        
        # Add headers
        header_cells = info_table.rows[0].cells
        header_cells[0].text = "Property"
        header_cells[1].text = "Value"
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Add basic file info
        row_cells = info_table.add_row().cells
        row_cells[0].text = "Original File"
        row_cells[1].text = os.path.basename(pim.rgb_path_original) if pim.rgb_path_original else "N/A"
        
        row_cells = info_table.add_row().cells
        row_cells[0].text = "Drone Model"
        row_cells[1].text = getattr(pim, "drone_model", "N/A")
        
        # Add EXIF information if available
        if hasattr(pim, 'exif') and pim.exif:
            exif_data = pim.exif
            
            # Common EXIF tags and their IDs
            exif_tags = {
                271: "Camera Make",            # Camera manufacturer
                272: "Camera Model",           # Camera model
                306: "Date/Time",             # Date and time
                36867: "Date/Time Original",  # Original date and time
                37377: "Shutter Speed",       # Shutter speed
                37378: "Aperture",            # Aperture
                37379: "Brightness",          # Brightness
                37380: "Exposure Comp.",      # Exposure bias
                37383: "Metering Mode",       # Metering mode
                37385: "Flash",               # Flash
                37386: "Focal Length"         # Focal length
            }
            
            for tag_id, tag_name in exif_tags.items():
                if tag_id in exif_data:
                    value = exif_data[tag_id]
                    # Format certain values
                    if tag_id == 37377 and value > 0:  # ShutterSpeed
                        value = f"1/{int(2**value)}"
                    elif tag_id == 37378:  # Aperture
                        value = f"f/{round(2**(value/2), 1)}"
                    elif tag_id == 37386:  # FocalLength
                        value = f"{value}mm"
                        
                    row_cells = info_table.add_row().cells
                    row_cells[0].text = tag_name
                    row_cells[1].text = str(value)
        
        # Add thermal parameters
        if hasattr(pim, 'thermal_param'):
            row_cells = info_table.add_row().cells
            row_cells[0].text = "Emissivity"
            row_cells[1].text = str(pim.thermal_param.get('emissivity', 'N/A'))
            
            row_cells = info_table.add_row().cells
            row_cells[0].text = "Distance (m)"
            row_cells[1].text = str(pim.thermal_param.get('distance', 'N/A'))
            
            row_cells = info_table.add_row().cells
            row_cells[0].text = "Humidity (%)"
            row_cells[1].text = str(pim.thermal_param.get('humidity', 'N/A'))
            
            row_cells = info_table.add_row().cells
            row_cells[0].text = "Reflection (°C)"
            row_cells[1].text = str(pim.thermal_param.get('reflection', 'N/A'))
        
        # Apply table styling based on template
        if isinstance(style_template, str):
            tpl_dict = STYLE_TEMPLATES.get(style_template, STYLE_TEMPLATES["modern_blue"])
            
            # Apply header styling
            if "table_header_bg" in tpl_dict and "table_header_text" in tpl_dict:
                header_bg = tpl_dict["table_header_bg"]
                header_text = tpl_dict["table_header_text"]
                
                # Apply to header row
                for cell in info_table.rows[0].cells:
                    # Set background color
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    cell_xml = cell._tc.get_or_add_tcPr()
                    cell_xml.append(parse_xml(
                        r'<w:shd {} w:fill="{:02X}{:02X}{:02X}"/>'.format(
                            nsdecls('w'), *header_bg)))
                    
                    # Set text color
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(*header_text)
            
            # Apply alternating row colors if specified
            if "table_row_alt" in tpl_dict:
                alt_color = tpl_dict["table_row_alt"]
                for i, row in enumerate(info_table.rows):
                    # Skip header row
                    if i == 0:
                        continue
                        
                    # Apply to alternating rows
                    if i % 2 == 1:  # Apply to odd rows (1-indexed, so 2nd, 4th, etc.)
                        for cell in row.cells:
                            from docx.oxml import parse_xml
                            from docx.oxml.ns import nsdecls
                            cell_xml = cell._tc.get_or_add_tcPr()
                            cell_xml.append(parse_xml(
                                r'<w:shd {} w:fill="{:02X}{:02X}{:02X}"/>'.format(
                                    nsdecls('w'), *alt_color)))
        
        # Add remarks if available
        if hasattr(pim, 'remarks') and pim.remarks and pim.remarks.strip():
            doc.add_heading('Remarks', level=3)
            remarks_para = doc.add_paragraph()
            remarks_para.add_run(pim.remarks)
        # Add key thermal information section
        doc.add_heading('Thermal Analysis', level=3)
        
        # Add temperature range information
        if hasattr(pim, 'tmin') and hasattr(pim, 'tmax'):
            temp_para = doc.add_paragraph()
            temp_para.add_run('Temperature Range: ').bold = True
            temp_para.add_run(f"{pim.tmin:.2f}°C to {pim.tmax:.2f}°C")
        
        # Add annotation tables if available with improved formatting
        if hasattr(pim, 'annot_rect_items') and pim.annot_rect_items:
            add_table_from_annotations(doc, pim.annot_rect_items, title="Rectangular Annotations", 
                                      style_template=tpl, annotation_type='rect')
        
        if hasattr(pim, 'meas_rect_list') and pim.meas_rect_list:
            add_table_from_annotations(doc, pim.meas_rect_list, title="Area Measurements", 
                                      style_template=tpl, annotation_type='rect')
        
        if hasattr(pim, 'meas_point_list') and pim.meas_point_list:
            add_table_from_annotations(doc, pim.meas_point_list, title="Spot Measurements", 
                                      style_template=tpl, annotation_type='point')
        
        if hasattr(pim, 'meas_line_list') and pim.meas_line_list:
            add_table_from_annotations(doc, pim.meas_line_list, title="Line Profile Measurements", 
                                      style_template=tpl, annotation_type='line')
        doc.add_page_break()

    doc.save(output_path)
    return output_path
